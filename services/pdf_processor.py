from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os

def extract_text_from_docx(file_path):
    doc=DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def process_document(file_path: str):
    ext = file_path.split('.')[-1].lower()
    source_name = os.path.basename(file_path)

    if ext == 'pdf':
        loader = PyPDFLoader(file_path)
        pages = loader.load()
    elif ext == 'docx':
        text = extract_text_from_docx(file_path)
        pages = [Document(page_content=text, metadata={"source": source_name})]
    else:
        raise ValueError("Unsupported file type")

    # Split pages or full text into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    chunks = splitter.split_documents(pages)

    # Add filename to metadata (in case not present)
    for chunk in chunks:
        chunk.metadata["source"] = source_name

    return chunks
